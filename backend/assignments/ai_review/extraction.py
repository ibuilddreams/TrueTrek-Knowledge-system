"""Extracts AI-readable content from assignment-related files.

Deliberately no new document-parsing subsystem: PDFs and images are sent to
Gemini natively as multimodal inline data (Gemini reads them directly,
including scanned/image-only PDFs — no separate OCR library needed since the
model itself reads the page images). Only `.docx` needs local text
extraction (via `python-docx` — pure-Python, no OS-level dependencies) since
Gemini's `generateContent` has no native Office-XML support. Legacy `.doc`
and `.ppt`/`.pptx`/`.zip` are already-allowed *upload* types (see
`assignments/validators.py`) but are not AI-readable without much heavier
dependencies — they are skipped for AI evaluation only; the file itself is
still stored and remains visible to teachers/admins as before.

`extract_file_content` is the shared, source-agnostic primitive — called
once over a submission's files (untrusted — the student's actual answer) and
once over an assignment's attachments (trusted — reference material/question
documents a teacher uploaded), via `extract_submission_content` and
`extract_reference_content` respectively. Keeping the same size/type/count
rules for both, budgeted independently, means a large reference PDF can
never crowd out the student's own submission in the same request.
"""

import base64
import os

from docx import Document as DocxDocument

from .exceptions import NoReadableContentError

# Gemini's native multimodal input handles these directly.
INLINE_MIME_TYPES = {
    ".pdf": "application/pdf",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".png": "image/png",
    ".webp": "image/webp",
}

# Extracted server-side since Gemini has no native Office-XML support.
DOCX_EXTENSIONS = {".docx"}

# Already plain text on disk — read and decoded directly, no parsing needed.
# Includes common programming-assignment source file types (e.g. a student
# submitting a bare .cpp file) alongside .txt/.md.
PLAIN_TEXT_EXTENSIONS = {
    ".txt", ".md",
    ".c", ".h", ".cpp", ".cc", ".cxx", ".hpp",
    ".py", ".java", ".js", ".ts", ".cs",
}

TEXT_EXTENSIONS = DOCX_EXTENSIONS | PLAIN_TEXT_EXTENSIONS

MAX_SINGLE_FILE_BYTES = 15 * 1024 * 1024
MAX_TOTAL_INLINE_BYTES = 15 * 1024 * 1024
MAX_EXTRACTED_TEXT_CHARS = 20_000
MAX_FILES_EVALUATED = 5


def _extract_docx_text(file_field):
    with file_field.open("rb") as fh:
        document = DocxDocument(fh)

    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    text = "\n".join(paragraphs).strip()
    if len(text) > MAX_EXTRACTED_TEXT_CHARS:
        text = text[:MAX_EXTRACTED_TEXT_CHARS] + "\n[...truncated...]"
    return text


def _extract_plain_text(file_field):
    with file_field.open("rb") as fh:
        raw = fh.read()
    text = raw.decode("utf-8", errors="replace").strip()
    if len(text) > MAX_EXTRACTED_TEXT_CHARS:
        text = text[:MAX_EXTRACTED_TEXT_CHARS] + "\n[...truncated...]"
    return text


def extract_file_content(files):
    """Given an iterable of objects each exposing `.file` (a Django
    FieldFile) and `.original_name`, returns (inline_files, text_blocks,
    unreadable_filenames) — the shared primitive behind both
    extract_submission_content (untrusted) and extract_reference_content
    (trusted). Never raises; an empty/all-unreadable input just returns
    empty lists — callers decide what "nothing readable" means for their
    context.

    inline_files:  [{"mime_type": str, "data_base64": str, "filename": str}]
    text_blocks:   [{"filename": str, "text": str}]
    unreadable_filenames: [str]
    """

    inline_files = []
    text_blocks = []
    unreadable_filenames = []
    inline_bytes_used = 0

    for item in files:
        filename = item.original_name or os.path.basename(item.file.name)
        extension = os.path.splitext(filename)[1].lower()

        if extension in TEXT_EXTENSIONS:
            try:
                if extension in DOCX_EXTENSIONS:
                    text = _extract_docx_text(item.file)
                else:
                    text = _extract_plain_text(item.file)
            except Exception:
                unreadable_filenames.append(filename)
                continue
            if text:
                text_blocks.append({"filename": filename, "text": text})
            else:
                unreadable_filenames.append(filename)
            continue

        if extension in INLINE_MIME_TYPES:
            if len(inline_files) >= MAX_FILES_EVALUATED:
                unreadable_filenames.append(filename)
                continue
            size = item.file.size
            if size > MAX_SINGLE_FILE_BYTES or inline_bytes_used + size > MAX_TOTAL_INLINE_BYTES:
                unreadable_filenames.append(filename)
                continue
            with item.file.open("rb") as fh:
                raw = fh.read()
            inline_files.append(
                {
                    "mime_type": INLINE_MIME_TYPES[extension],
                    "data_base64": base64.b64encode(raw).decode("ascii"),
                    "filename": filename,
                }
            )
            inline_bytes_used += size
            continue

        unreadable_filenames.append(filename)

    return inline_files, text_blocks, unreadable_filenames


def extract_submission_content(submission):
    """Returns (inline_files, text_blocks, unreadable_filenames) for the
    student's files uploaded in the *current* submission attempt — untrusted
    content to evaluate.

    `assignments.services.submit_assignment` never deletes prior files on
    resubmission (existing, unmodified behavior — files are additive), so
    "current attempt" is derived here by filtering on
    `created_at >= submission.submitted_at` rather than changing that shared
    function's behavior for MANUAL/AUTO assignments.

    Raises NoReadableContentError if nothing usable was found at all.
    """

    files = list(submission.files.filter(created_at__gte=submission.submitted_at))
    if not files:
        # Defensive fallback (e.g. a directly-constructed submission in a
        # test) — never silently evaluate zero files when some exist.
        files = list(submission.files.all())

    inline_files, text_blocks, unreadable_filenames = extract_file_content(files)

    if not inline_files and not text_blocks:
        raise NoReadableContentError("None of the submitted files could be read for AI review.")

    return inline_files, text_blocks, unreadable_filenames


def extract_reference_content(assignment):
    """Returns (inline_files, text_blocks, unreadable_filenames) for the
    assignment's teacher/admin-uploaded attachments — trusted reference
    material (e.g. a PDF defining the actual questions/instructions),
    budgeted independently of the student's submission. Never raises: a
    reference-free assignment (the common case — most AI-graded assignments
    only need `description`) just returns empty lists."""

    attachments = list(assignment.attachments.all())
    if not attachments:
        return [], [], []

    return extract_file_content(attachments)
